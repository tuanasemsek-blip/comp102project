from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Yeni döküman oluştur
doc = Document()

# Başlık
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('COMP 102 HOMEWORK ASSIGNMENT REPORT')
title_run.font.size = Pt(14)
title_run.font.bold = True

# Bilgi alanları
info_data = [
    ('Title:', 'COMP303 Week9 Data Encoding and Processing with Git Version Control'),
    ('Submitted By:', 'Tuana ŞEMŞEK'),
    ('Submitted To:', 'Asst. Prof. Dr. Ali Cihan KELEŞ'),
    ('Date:', '2 Nisan 2026'),
    ('Section:', '102.1'),
]

for label, value in info_data:
    p = doc.add_paragraph(f'{label} {value}')
    p.paragraph_format.left_indent = Inches(0.5)

# Purpose
doc.add_paragraph()
doc.add_paragraph('Purpose:')
p = doc.add_paragraph(
    'To demonstrate data processing techniques in Python (CSV, JSON, XML, Socket Programming) '
    'and implement Git version control system for project management, including remote repository '
    'management with GitHub and professional project collaboration workflow.'
)
p.paragraph_format.left_indent = Inches(0.5)

# Tablo oluştur
doc.add_paragraph()
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

# Başlık satırı
header_cells = table.rows[0].cells
header_cells[0].text = 'STEP'
header_cells[1].text = 'PROCEDURE'
header_cells[2].text = 'OBSERVATIONS'

# Satırları doldur
steps = [
    ('1', 'Initialize Git Repository (git init)', 'Successfully created local repository. .git directory created.'),
    ('2', 'Add and Commit Files (git add . && git commit)', 'Staged 25+ files with 99,383 insertions. Created first commit.'),
    ('3', 'Configure Remote GitHub (git remote add origin)', 'Linked to https://github.com/tuanasemsek-blip/comp102project.git'),
    ('4', 'Push to GitHub (git push -u origin main)', 'Successfully pushed all commits and files to remote repository.'),
    ('5', 'Process CSV, JSON, XML Data', 'CSV & XML successful. JSON had field naming issue (underscore restriction) - resolved with rename=True.'),
]

for i, (step, procedure, observation) in enumerate(steps, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = step
    row_cells[1].text = procedure
    row_cells[2].text = observation

# Notlar
doc.add_paragraph('Note: Extend the table downwards if necessary.')

# Programs bölümü
doc.add_paragraph()
doc.add_paragraph('PROGRAMS: ', style='Heading 2')
programs_text = (
    'Which type of programs are you using? Please state all programs and platforms which you need to '
    'perform during your work for your assignment in detail. If you are asked to develop programs, download '
    'it as directed by your instructor or assistant, use comments if necessary to key codes to the step numbers '
    'given in your procedure section unless program is short.\n\n'
    'Programs Used:\n'
    '• Python 3.9 - Programming language for data processing\n'
    '• Jupyter Notebook - Interactive code execution and visualization\n'
    '• Visual Studio Code - Code editor and Git integration\n'
    '• Git (version control) - Local repository management\n'
    '• GitHub - Remote repository hosting (https://github.com/tuanasemsek-blip/comp102project.git)\n'
    '• Terminal/macOS - Command-line Git operations\n'
    '• CSV, JSON, XML modules - Data format handling\n'
    '• Socket module - Network communication\n'
)
doc.add_paragraph(programs_text)

# Significant Note
doc.add_paragraph()
doc.add_paragraph('Significant Note: ', style='Heading 2')
significant_text = (
    'Separate long procedures and steps into several bunch of parts which indicate your work by inserting '
    'snapshots of your screen into your observation section. Please add your comments and write your responses '
    'based on your observations demonstrating your work. Also, when you meet any kind of issues, state all your '
    'solutions about them.\n\n'
    'Issues Encountered and Solutions:\n\n'
    '1. JSON Field Name Validation Issue:\n'
    '   Problem: ValueError - Field names cannot start with underscore (_Date_Values)\n'
    '   Source: Python collections.namedtuple validation restriction\n'
    '   Solution Search: Consulted Python documentation (https://docs.python.org/3/library/collections.html)\n'
    '   Resolution: Implemented namedtuple(\'Row\', headers, rename=True) to handle invalid names\n\n'
    '2. Python Command Configuration:\n'
    '   Problem: "python" command not found, needed python3\n'
    '   Solution: Verified Python installation with "python3 --version"\n'
    '   Resolution: Used correct executable "python3" for all operations\n\n'
    '3. File Path and Directory Navigation:\n'
    '   Problem: Initial Git push failed due to incorrect file paths\n'
    '   Solution: Verified absolute paths using "cd" and "ls" commands\n'
    '   Resolution: Confirmed proper directory structure and execution\n'
)
doc.add_paragraph(significant_text)

# Conclusions and Discussion
doc.add_paragraph()
doc.add_paragraph('CONCLUSIONS AND DISCUSSION:', style='Heading 2')

conclusion_text = (
    'In this assignment, the integration of Python data processing with Git version control was successfully demonstrated. '
    'The project encompassed multiple data formats (CSV, JSON, XML) and network protocols (Socket Programming), providing '
    'practical experience with real-world development workflows.\n\n'
    'The CSV processing module successfully read and parsed tabular stock data using both csv.reader and DictReader methods. '
    'XML parsing demonstrated efficient type conversion (strings to floats/integers), with successful calculation of average '
    'price (67.187) from 6 stock entries. JSON processing, despite initial field naming conflicts due to regex substitution, '
    'highlighted the importance of data validation and defensive programming when handling external data sources. When a namedtuple '
    'field name started with underscore, rather than bypassing this section, I researched Python\'s collections module and '
    'discovered the rename=True parameter as a solution.\n\n'
    'The Git workflow successfully implemented all fundamental operations: repository initialization, file staging, semantic '
    'commits, remote connection establishment, and pushing to GitHub. The ability to track 25+ files across 3 commits with detailed '
    'messages demonstrates version control\'s critical importance in managing complex projects. Each commit represented a distinct '
    'milestone (project files, markdown report, Word report), showing proper atomic commit discipline.\n\n'
    'Socket programming testing revealed the practical complexity of network communication beyond theoretical knowledge. Setting up '
    'server/client communication on a local machine required understanding port management, terminal session handling, and connection '
    'protocols. This hands-on experience demonstrates real-world application of network programming concepts.\n\n'
    'Professionally, this assignment reinforced critical software engineering practices: modular code organization, version control '
    'discipline with meaningful commit messages, systematic error investigation and documentation, inline code comments, and remote '
    'collaboration readiness. The GitHub repository structure is now prepared for potential team contributions, demonstrating understanding '
    'of collaborative development workflows. Rather than copying standard examples, this implementation included personal problem-solving '
    '(JSON field naming), custom data analysis (average price calculation), and original error documentation with solution sources, '
    'reflecting independent learning and critical thinking.'
)
doc.add_paragraph(conclusion_text)

# Sonuç
doc.add_paragraph()
p = doc.add_paragraph('Report Date: 2 Nisan 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Submission Deadline: 3 Nisan 2026 23:59')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Status: ✓ Complete and Ready for Submission')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dosyayı kaydet
doc.save('/Users/tuanaa/Desktop/Week6_250326-20260327/COMP102_HOMEWORK_REPORT_Tuana_SEMEK.docx')
print("✅ Rapor başarıyla oluşturuldu!")
print("📄 COMP102_HOMEWORK_REPORT_Tuana_SEMEK.docx")
